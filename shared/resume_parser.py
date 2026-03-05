import fitz  # PyMuPDF
import docx
import os
import tempfile


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes."""
    text = ""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        doc = fitz.open(tmp_path)
        for page in doc:
            text += page.get_text()
        doc.close()
    finally:
        os.unlink(tmp_path)

    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    text = ""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        doc = docx.Document(tmp_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs)
    finally:
        os.unlink(tmp_path)

    return text.strip()


def parse_resume(uploaded_file) -> tuple[str, str]:
    """
    Parse resume from Streamlit UploadedFile object.
    Returns: (extracted_text, file_extension)
    """
    file_bytes = uploaded_file.read()
    file_name = uploaded_file.name.lower()

    if file_name.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
        return text, "pdf"
    elif file_name.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
        return text, "docx"
    else:
        raise ValueError("Unsupported file format. Please upload PDF or DOCX.")


def save_resume_file(file_bytes: bytes, user_id: int, ext: str, upload_dir: str = "uploads") -> str:
    """Save resume file to disk and return path."""
    os.makedirs(upload_dir, exist_ok=True)
    file_path = os.path.join(upload_dir, f"resume_{user_id}.{ext}")
    with open(file_path, "wb") as f:
        f.write(file_bytes)
    return file_path
